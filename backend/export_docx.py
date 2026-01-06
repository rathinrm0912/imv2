from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any, List
import json

def generate_docx(doc_data: Dict[str, Any], output_path: str):
    """Generate a formatted Word document from IM data"""
    document = Document()
    
    # Set document styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'IBM Plex Sans'
    font.size = Pt(11)
    
    # Title
    title = document.add_heading('Redwood Partners Investment Memorandum', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Document metadata
    info = document.add_paragraph()
    info.add_run(f"Title: {doc_data['title']}\n").bold = True
    info.add_run(f"Status: {doc_data['status'].upper()}\n")
    info.add_run(f"Created: {doc_data['created_at']}\n")
    info.add_run(f"Last Updated: {doc_data['updated_at']}\n")
    
    document.add_paragraph()  # Spacing
    document.add_page_break()
    
    # Table of Contents
    toc_heading = document.add_heading('Table of Contents', 1)
    for section in doc_data.get('sections', []):
        if not section.get('show_instructions', True):  # Skip instructions if hidden
            toc_entry = document.add_paragraph()
            toc_entry.add_run(f"{section['section_number']}. {section['title']}")
            toc_entry.paragraph_format.left_indent = Inches(0.25)
    
    document.add_page_break()
    
    # Sections
    for section in doc_data.get('sections', []):
        # Section heading
        heading = document.add_heading(f"{section['section_number']}. {section['title']}", 1)
        heading.runs[0].font.color.rgb = RGBColor(6, 78, 59)  # Redwood green
        
        # Only show instructions if enabled
        if section.get('show_instructions', True) and section.get('instructions'):
            instructions = document.add_paragraph()
            instructions.add_run('Instructions: ').bold = True
            instructions.add_run(section['instructions'])
            instructions.runs[1].font.italic = True
            instructions.runs[1].font.color.rgb = RGBColor(100, 116, 139)  # Muted
        
        # Section content
        content = section.get('content', {})
        if content:
            for key, value in content.items():
                if value:
                    field_para = document.add_paragraph()
                    field_para.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                    field_para.add_run(str(value))
        else:
            document.add_paragraph('[Content to be added]')
        
        document.add_paragraph()  # Spacing between sections
    
    # Save document
    document.save(output_path)
    return output_path
